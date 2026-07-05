"""Document export tools: PDF and DOCX generation with proper Markdown parsing."""
import os
import re
import logging
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document

logger = logging.getLogger(__name__)


def _escape_xml(text: str) -> str:
    """Escape special XML characters for ReportLab."""
    if not text:
        return ""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    # Handle markdown bold/italic after escaping
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Handle inline code
    text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
    # Handle links - convert [text](url) to clickable
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2" color="blue">\1</a>', text)
    return text


def clean_text(text: str) -> str:
    """Clean text before document generation."""
    if not text:
        return "No content generated."
    text = text.replace('<para>', '').replace('</para>', '')
    text = text.replace('<br>', '\n')
    return text.strip()


def save_as_pdf(markdown_text: str, filename: str):
    """Generate PDF with proper Markdown parsing and citation support."""
    os.makedirs(os.path.dirname(filename) if os.path.dirname(filename) else ".", exist_ok=True)

    cleaned_text = clean_text(markdown_text)

    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'],
        fontSize=20, spaceAfter=8, textColor=colors.HexColor('#1a1a2e')
    )
    subtitle_style = ParagraphStyle(
        'Subtitle', parent=styles['Normal'],
        fontSize=10, textColor=colors.HexColor('#666666'), spaceAfter=20
    )
    h1_style = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        fontSize=16, spaceBefore=16, spaceAfter=8, textColor=colors.HexColor('#1a1a2e')
    )
    h2_style = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        fontSize=13, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor('#302b63')
    )
    h3_style = ParagraphStyle(
        'H3', parent=styles['Heading3'],
        fontSize=11, spaceBefore=10, spaceAfter=4, textColor=colors.HexColor('#24243e')
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontSize=10, leading=14, spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'Bullet', parent=body_style,
        leftIndent=20, bulletIndent=10, spaceAfter=4
    )
    ref_style = ParagraphStyle(
        'Reference', parent=body_style,
        fontSize=9, textColor=colors.HexColor('#555555'), leftIndent=15
    )

    # Title page
    story.append(Paragraph("ResearchPilot AI — Technical Report", title_style))
    story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), subtitle_style))
    story.append(Spacer(1, 20))

    # Process content
    in_references = False
    lines = cleaned_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect References section
        if line.startswith('## References') or line.startswith('# References'):
            in_references = True
            story.append(Spacer(1, 10))
            story.append(Paragraph(_escape_xml(line.lstrip('#').strip()), h1_style))
            continue

        # Skip horizontal rules
        if line.startswith('---'):
            story.append(Spacer(1, 8))
            continue

        # Headings
        if line.startswith('# '):
            story.append(Paragraph(_escape_xml(line[2:]), h1_style))
        elif line.startswith('## '):
            story.append(Paragraph(_escape_xml(line[3:]), h2_style))
        elif line.startswith('### '):
            story.append(Paragraph(_escape_xml(line[4:]), h3_style))
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            prefix = line[:2] if line[0] in '-*' else line[:2]
            story.append(Paragraph(f"• {_escape_xml(line[2:].strip())}", bullet_style))
        elif in_references and line.startswith('['):
            story.append(Paragraph(_escape_xml(line), ref_style))
        else:
            style = ref_style if in_references else body_style
            story.append(Paragraph(_escape_xml(line), style))

    try:
        doc.build(story)
        logger.info(f"PDF saved: {filename}")
    except Exception as e:
        logger.error(f"PDF build failed: {e}")
        # Fallback: minimal PDF
        doc2 = SimpleDocTemplate(filename, pagesize=letter)
        doc2.build([Paragraph("Report generation encountered formatting issues. Please use Markdown or DOCX export.", styles['Normal'])])

    return filename


def save_as_docx(markdown_text: str, filename: str):
    """Generate DOCX with proper Markdown parsing and citation support."""
    os.makedirs(os.path.dirname(filename) if os.path.dirname(filename) else ".", exist_ok=True)

    cleaned_text = clean_text(markdown_text)
    doc = Document()

    # Title
    title = doc.add_heading('ResearchPilot AI — Technical Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    doc.add_paragraph("─" * 50)

    # Process content
    lines = cleaned_text.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = 80000  # ~8pt in EMU
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip horizontal rules
        if stripped.startswith('---'):
            doc.add_paragraph("─" * 50)
            continue

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
            doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.', stripped):
            doc.add_paragraph(stripped, style='List Number')
        else:
            p = doc.add_paragraph()
            # Handle bold and italic in text
            _add_formatted_text(p, stripped)

    doc.save(filename)
    logger.info(f"DOCX saved: {filename}")
    return filename


def _add_formatted_text(paragraph, text: str):
    """Add text with bold/italic formatting to a DOCX paragraph."""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic within non-bold parts
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*'):
                    run = paragraph.add_run(ipart[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ipart)
